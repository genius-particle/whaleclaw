"""DOCX edit tool — text replacement + image operations in an existing .docx file."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from whaleclaw.tools.base import Tool, ToolDefinition, ToolParameter, ToolResult
from whaleclaw.tools.deps import ensure_tool_dep


class DocxEditTool(Tool):
    """Edit an existing Word document: replace text, add/replace/remove images."""

    @property
    def definition(self) -> ToolDefinition:
        return ToolDefinition(
            name="docx_edit",
            description=(
                "修改现有 Word（.docx）文档：支持文本替换、插入图片、替换已有图片、删除图片。"
                "换图请用 replace_image（保持原位置尺寸），新增图片用 add_image。"
            ),
            parameters=[
                ToolParameter(name="path", type="string", description="DOCX 文件绝对路径"),
                ToolParameter(
                    name="action",
                    type="string",
                    description=(
                        "操作类型：replace_text|add_image|replace_image|remove_image"
                    ),
                    required=False,
                    enum=[
                        "replace_text",
                        "add_image",
                        "replace_image",
                        "remove_image",
                    ],
                ),
                ToolParameter(
                    name="old_text",
                    type="string",
                    description="要替换的原文（replace_text 时必填）",
                    required=False,
                ),
                ToolParameter(
                    name="new_text",
                    type="string",
                    description="替换后的新文案（replace_text 时必填）",
                    required=False,
                ),
                ToolParameter(
                    name="image_path",
                    type="string",
                    description="图片绝对路径（add_image/replace_image 时必填）",
                    required=False,
                ),
                ToolParameter(
                    name="image_index",
                    type="integer",
                    description="目标图片序号（从 1 开始，replace_image/remove_image 时使用，默认 1）",
                    required=False,
                ),
                ToolParameter(
                    name="image_width_inches",
                    type="number",
                    description="图片宽度（英寸，add_image 时可选，默认 5.0）",
                    required=False,
                ),
                ToolParameter(
                    name="paragraph_index",
                    type="integer",
                    description="在第 N 段后插入图片（从 1 开始，add_image 时可选，默认追加到末尾）",
                    required=False,
                ),
            ],
        )

    async def execute(self, **kwargs: Any) -> ToolResult:
        if not ensure_tool_dep("docx"):
            return ToolResult(success=False, output="", error="缺少依赖 python-docx")

        from docx import Document
        from docx.shared import Inches

        raw_path = str(kwargs.get("path", "")).strip()
        action = str(kwargs.get("action", "replace_text")).strip().lower() or "replace_text"
        old_text = str(kwargs.get("old_text", "")).strip()
        new_text = str(kwargs.get("new_text", ""))
        image_path = str(kwargs.get("image_path", "")).strip()

        if not raw_path:
            return ToolResult(success=False, output="", error="path 不能为空")

        path = Path(raw_path).expanduser().resolve()
        if not path.is_file():
            return ToolResult(success=False, output="", error=f"文件不存在: {path}")
        if path.suffix.lower() != ".docx":
            return ToolResult(success=False, output="", error="仅支持 .docx 文件")

        try:
            doc = Document(str(path))
        except Exception as exc:
            return ToolResult(success=False, output="", error=f"DOCX 打开失败: {exc}")

        if action == "replace_text":
            if not old_text:
                return ToolResult(success=False, output="", error="old_text 不能为空")
            replaced = _replace_text_in_doc(doc, old_text, new_text)
            if replaced == 0:
                return ToolResult(success=False, output="", error="文档中未找到匹配文本")
            output = f"已修改 {path}，替换 {replaced} 处文本"

        elif action == "add_image":
            if not image_path:
                return ToolResult(success=False, output="", error="add_image 需要 image_path")
            img = Path(image_path).expanduser().resolve()
            if not img.is_file():
                return ToolResult(success=False, output="", error=f"图片不存在: {img}")
            width = Inches(float(kwargs.get("image_width_inches", 5.0)))
            raw_para_idx = kwargs.get("paragraph_index")
            if raw_para_idx is not None:
                try:
                    para_idx = int(raw_para_idx)
                except (TypeError, ValueError):
                    para_idx = 0
                if para_idx < 1 or para_idx > len(doc.paragraphs):
                    return ToolResult(
                        success=False, output="",
                        error=f"paragraph_index={para_idx} 越界，共 {len(doc.paragraphs)} 段",
                    )
                target_para = doc.paragraphs[para_idx - 1]
                run = target_para.add_run()
                run.add_picture(str(img), width=width)
            else:
                doc.add_picture(str(img), width=width)
            output = f"已在 {path} 中插入图片 {img.name}"

        elif action in ("replace_image", "remove_image"):
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            image_rels = _collect_image_rels(doc)
            if not image_rels:
                return ToolResult(
                    success=False, output="",
                    error="文档中没有图片可操作",
                )
            try:
                img_idx = int(kwargs.get("image_index", 1))
            except (TypeError, ValueError):
                img_idx = 1
            if img_idx < 1 or img_idx > len(image_rels):
                return ToolResult(
                    success=False, output="",
                    error=f"image_index={img_idx} 越界，文档共 {len(image_rels)} 张图片",
                )

            target_blip, target_rel = image_rels[img_idx - 1]

            if action == "remove_image":
                _remove_drawing_element(target_blip)
                output = f"已删除 {path} 中第 {img_idx} 张图片"
            else:
                if not image_path:
                    return ToolResult(
                        success=False, output="",
                        error="replace_image 需要 image_path",
                    )
                img = Path(image_path).expanduser().resolve()
                if not img.is_file():
                    return ToolResult(success=False, output="", error=f"图片不存在: {img}")

                _replace_image_blob(doc, target_rel, img, RT)
                output = f"已替换 {path} 中第 {img_idx} 张图片为 {img.name}"
        else:
            return ToolResult(success=False, output="", error=f"不支持的 action: {action}")

        try:
            doc.save(str(path))
        except Exception as exc:
            return ToolResult(success=False, output="", error=f"DOCX 保存失败: {exc}")

        return ToolResult(success=True, output=output)


def _replace_text_in_doc(doc: Any, old_text: str, new_text: str) -> int:
    replaced = 0
    for para in doc.paragraphs:
        text = para.text or ""
        count = text.count(old_text)
        if count <= 0:
            continue
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced += count
                break
        else:
            para.text = text.replace(old_text, new_text)
            replaced += count

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                count = text.count(old_text)
                if count <= 0:
                    continue
                for para in cell.paragraphs:
                    for run in para.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, new_text)
                            replaced += count
                            break
                    else:
                        continue
                    break
                else:
                    cell.text = text.replace(old_text, new_text)
                    replaced += count
    return replaced


def _collect_image_rels(doc: Any) -> list[tuple[Any, Any]]:
    """Collect (blip_element, relationship) pairs for all images in document order."""
    from lxml import etree

    nsmap = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    results: list[tuple[Any, Any]] = []
    body = doc.element.body
    for blip in body.iter(etree.QName(nsmap["a"], "blip").text):  # type: ignore[arg-type]
        r_embed = blip.get(etree.QName(nsmap["r"], "embed").text)  # type: ignore[arg-type]
        if r_embed is None:
            continue
        try:
            rel = doc.part.rels[r_embed]
        except KeyError:
            continue
        results.append((blip, rel))
    return results


def _remove_drawing_element(blip: Any) -> None:
    """Walk up from a:blip to the nearest w:drawing or w:pict and remove it."""
    node = blip
    while node is not None:
        tag = node.tag
        if isinstance(tag, str) and (tag.endswith("}drawing") or tag.endswith("}pict")):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)
            return
        node = node.getparent()


def _replace_image_blob(doc: Any, rel: Any, new_img_path: Path, rt: Any) -> None:
    """Replace the image data behind a relationship with a new file."""
    import mimetypes

    from docx.opc.part import Part

    content_type = mimetypes.guess_type(str(new_img_path))[0] or "image/png"
    blob = new_img_path.read_bytes()

    old_part = rel.target_part
    new_part = Part(
        old_part.partname,
        content_type,
        blob,
        old_part.package,
    )
    rel._target = new_part  # pyright: ignore[reportPrivateUsage]
